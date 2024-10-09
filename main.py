# coding: utf8
import pandas as pd
import numpy as np
import docx

table = pd.read_excel('2024-10-01 Konferentsiia actual.xlsx')

# список с названиями секций
sections_list = ['СЕКЦИЯ ФИЛОСОФИИ', 'СЕКЦИЯ СОЦИОЛОГИИ НАУКИ И ТЕХНОЛОГИЙ', 'СЕКЦИЯ ИСТОРИИ', 'СЕКЦИЯ ИНОСТРАННЫХ ЯЗЫКОВ И РУССКОГО КАК ИНОСТРАННОГО',
                 'СЕКЦИЯ ПРОМЫШЛЕННОЙ ЭКОЛОГИИ И БЕЗОПАСНОСТИ ЖИЗНЕДЕЯТЕЛЬНОСТИ', 'СЕКЦИЯ ИНЖЕНЕРНОЙ ГРАФИКИ', 'СЕКЦИЯ ХИМИИ', 'СЕКЦИЯ ВЫСШЕЙ МАТЕМАТИКИ',
                 'СЕКЦИЯ ОБЩЕЙ ИНФОРМАТИКИ', 'СЕКЦИЯ НАНОЭЛЕКТРОНИКИ И МИКРОСИСТЕМНОЙ ТЕХНИКИ', 'СЕКЦИЯ ТЕОРЕТИЧЕСКОЙ МЕХАНИКИ', 'СЕКЦИЯ СОПРОТИВЛЕНИЯ МАТЕРИАЛОВ',
                 'СЕКЦИЯ КОСМИЧЕСКОГО МАШИНОСТРОЕНИЯ', 'СЕКЦИЯ МОЛОДЕЖНОЙ АЭРОКОСМИЧЕСКОЙ ШКОЛЫ', 'СЕКЦИЯ КОСМИЧЕСКИХ ИССЛЕДОВАНИЙ И НАНОСПУТНИКОВЫХ ТЕХНОЛОГИЙ', 'СЕКЦИЯ ТЕОРИИ ДВИГАТЕЛЕЙ ЛЕТАТЕЛЬНЫХ АППАРАТОВ',
                 'СЕКЦИЯ КОНСТРУКЦИИ И ПРОЕКТИРОВАНИЯ ДВИГАТЕЛЕЙ ЛЕТАТЕЛЬНЫХ АППАРАТОВ', 'СЕКЦИЯ ТЕХНОЛОГИЙ ПРОИЗВОДСТВА ДВИГАТЕЛЕЙ', 'СЕКЦИЯ АВТОМАТИЧЕСКИХ СИСТЕМ ЭНЕРГЕТИЧЕСКИХ УСТАНОВОК', 'СЕКЦИЯ ТЕПЛОТЕХНИКИ И ТЕПЛОВЫХ ДВИГАТЕЛЕЙ',
                 'СЕКЦИЯ ДИНАМИКИ ПОЛЁТА И СИСТЕМ УПРАВЛЕНИЯ', 'СЕКЦИЯ КОНСТРУКЦИИ И ПРОЕКТИРОВАНИЯ ЛЕТАТЕЛЬНЫХ АППАРАТОВ', 'СЕКЦИЯ ПРОИЗВОДСТВА ЛЕТАТЕЛЬНЫХ АППАРАТОВ И УПРАВЛЕНИЯ КАЧЕСТВОМ В МАШИНОСТРОЕНИИ',
                 'СЕКЦИЯ ЭКСПЛУАТАЦИИ АВИАЦИОННОЙ ТЕХНИКИ', 'СЕКЦИЯ ОСНОВ КОНСТРУИРОВАНИЯ МАШИН', 'СЕКЦИЯ ОРГАНИЗАЦИИ ПЕРЕВОЗОК И УПРАВЛЕНИЯ НА ТРАНСПОРТЕ',
                 'СЕКЦИЯ ВОЕННОГО УЧЕБНОГО ЦЕНТРА', 'СЕКЦИЯ ТЕХНОЛОГИИ МЕТАЛЛОВ И АВИАМАТЕРИАЛОВЕДЕНИЯ', 'СЕКЦИЯ ТЕХНОЛОГИИ МЕТАЛЛУРГИИ И МАШИНОСТРОЕНИЯ', 'СЕКЦИЯ ЭЛЕКТРОТЕХНИКИ',
                 'СЕКЦИЯ РАДИОТЕХНИЧЕСКИХ ИНФОРМАЦИОННЫХ СИСТЕМ', 'СЕКЦИЯ БИОМЕДИЦИНСКОЙ ИНЖЕНЕРИИ И БИОФОТОНИКИ', 'СЕКЦИЯ КОНСТРУИРОВАНИЯ И ТЕХНОЛОГИЙ ЭЛЕКТРОННЫХ СРЕДСТВ И УСТРОЙСТВ',
                 'СЕКЦИЯ ОБЩЕЙ И ТЕОРЕТИЧЕСКОЙ ФИЗИКИ', 'СЕКЦИЯ ФИЗИКИ ТВЕРДОГО ТЕЛА И ПОЛУПРОВОДНИКОВОЙ НАНОЭЛЕКТРОНИКИ', 'СЕКЦИЯ ОПТИКИ И СПЕКТРОСКОПИИ', 'СЕКЦИЯ ТЕХНИЧЕСКОЙ КИБЕРНЕТИКИ',
                 'СЕКЦИЯ ПРИКЛАДНЫХ МАТЕМАТИКИ И ФИЗИКИ', 'СЕКЦИЯ ГЕОИНФОРМАТИКИ, ОБРАБОТКИ ИЗОБРАЖЕНИЙ И ИНФОРМАЦИОННОЙ БЕЗОПАСНОСТИ', 'СЕКЦИЯ ИНФОРМАЦИОННЫХ СИСТЕМ И ТЕХНОЛОГИЙ',
                 'СЕКЦИЯ ПРОГРАММНЫХ СИСТЕМ', 'СЕКЦИЯ АЛГЕБРЫ И ГЕОМЕТРИИ', 'СЕКЦИЯ ФУНКЦИОНАЛЬНОГО АНАЛИЗА И ТЕОРИИ ФУНКЦИЙ', 'СЕКЦИЯ ИНФОРМАТИКИ И ВЫЧИСЛИТЕЛЬНОЙ МАТЕМАТИКИ', 'СЕКЦИЯ ДИФФЕРЕНЦИАЛЬНЫХ УРАВНЕНИЙ И ТЕОРИИ УПРАВЛЕНИЯ',
                 'СЕКЦИЯ БЕЗОПАСНОСТИ ИНФОРМАЦИОННЫХ СИСТЕМ', 'СЕКЦИЯ МАТЕМАТИЧЕСКОГО МОДЕЛИРОВАНИЯ В МЕХАНИКЕ', 'СЕКЦИЯ ОРГАНИЧЕСКОЙ ХИМИИ', 'СЕКЦИЯ ХИМИИ ПРИРОДНЫХ И ВЫСОКОМОЛЕКУЛЯРНЫХ СОЕДИНЕНИЙ',
                 'СЕКЦИЯ ХИМИИ ГЕТЕРОЦИКЛИЧЕСКИХ СОЕДИНЕНИЙ', 'СЕКЦИЯ НЕОРГАНИЧЕСКОЙ ХИМИИ', 'СЕКЦИЯ ФИЗИЧЕСКОЙ, АНАЛИТИЧЕСКОЙ ХИМИИ И ХРОМАТОГРАФИИ', 'СЕКЦИЯ БИОХИМИИ, БИОТЕХНОЛОГИИ И БИОИНЖЕНЕРИИ',
                 'СЕКЦИЯ ФИЗИОЛОГИИ ЧЕЛОВЕКА И ЖИВОТНЫХ', 'СЕКЦИЯ СОВРЕМЕННЫХ ПРОБЛЕМ АУТЭКОЛОГИИ', 'СЕКЦИЯ РЕГИОНАЛЬНЫХ ПРОБЛЕМ ОХРАНЫ ПРИРОДЫ', 'СЕКЦИЯ ЭКОЛОГИИ ПОЧВ И ВОДНОЙ СРЕДЫ',
                 'СЕКЦИЯ МЕНЕДЖМЕНТА И ОРГАНИЗАЦИИ ПРОИЗВОДСТВА', 'СЕКЦИЯ МАТЕМАТИЧЕСКИХ МЕТОДОВ ТЕОРИИ УПРАВЛЕНИЯ', 'СЕКЦИЯ БИЗНЕС-ИНФОРМАТИКИ', 'СЕКЦИЯ АКТУАЛЬНЫХ ПРАВОВЫХ И СОЦИАЛЬНЫХ ИССЛЕДОВАНИЙ',
                 'СЕКЦИЯ ЭКОНОМИКИ И ФИНАНСОВ', 'СЕКЦИЯ «ФИНАНСОВО-ЭКОНОМИЧЕСКАЯ АНАЛИТИКА В СИСТЕМЕ УПРАВЛЕНИЯ СОЦИАЛЬНО-ЭКОНОМИЧЕСКИМИ И БИЗНЕС-ПРОЦЕССАМИ»',
                 'СЕКЦИЯ «ГЛОБАЛЬНАЯ ТРАНСФОРМАЦИЯ XXI ВЕКА: ПЕРЕХОД В НОВУЮ ЭРУ ПОТРЕБЛЕНИЯ ЧЕЛОВЕЧЕСТВА»', 'СЕКЦИЯ АКТУАЛЬНЫХ ПРОБЛЕМ ФУНКЦИОНИРОВАНИЯ РЕГИОНАЛЬНЫХ СОЦИАЛЬНО-ЭКОНОМИЧЕСКИХ СИСТЕМ И МИРОВОЙ ЭКОНОМИКИ',
                 'СЕКЦИЯ ЦИФРОВИЗАЦИИ, ИННОВАЦИОННЫХ И ФИНАНСОВЫХ АСПЕКТОВ ФУНЦИОНИРОВАНИЯ СОВРЕМЕННОЙ ЭКОНОМИКИ', 'СЕКЦИЯ ОБЩЕГО И СТРАТЕГИЧЕСКОГО МЕНЕДЖМЕНТА', 'СЕКЦИЯ УПРАВЛЕНИЯ ПРОЕКТАМИ',
                 'СЕКЦИЯ МЕНЕДЖМЕНТА ПЕРСОНАЛА', 'СЕКЦИЯ МАРКЕТИНГА', 'СЕКЦИЯ ФИНАНСОВОГО МЕНЕДЖМЕНТА', 'СЕКЦИЯ СВЯЗЕЙ С ОБЩЕСТВЕННОСТЬЮ В ОРГАНАХ ГОСУДАРСТВЕННОЙ ВЛАСТИ',
                 'СЕКЦИЯ ТЕОРИИ И ПРАКТИКИ ГОСУДАРСТВЕННОГО УПРАВЛЕНИЯ', 'СЕКЦИЯ АНАЛИЗА И ПРОГРАММИРОВАНИЯ РАЗВИТИЯ СОЦИАЛЬНОЙ СФЕРЫ', 'СЕКЦИЯ ПРЕДПРИНИМАТЕЛЬСТВА И ГОСУДАРСТВЕННОЙ ВЛАСТИ',
                 'СЕКЦИЯ УПРАВЛЕНИЯ ПЕРСОНАЛОМ', 'СЕКЦИЯ ТЕОРИИ ГОСУДАРСТВА И ПРАВА', 'СЕКЦИЯ ИСТОРИИ ГОСУДАРСТВА И ПРАВА РОССИИ', 'СЕКЦИЯ ИСТОРИИ ГОСУДАРСТВА И ПРАВА ЗАРУБЕЖНЫХ СТРАН',
                 'СЕКЦИЯ МЕЖДУНАРОДНОГО ПУБЛИЧНОГО ПРАВА', 'СЕКЦИЯ ПРОБЛЕМ КОСМИЧЕСКОГО ПРАВА', 'СЕКЦИЯ РИМСКОГО ПРАВА', 'СЕКЦИЯ АКТУАЛЬНЫХ ПРОБЛЕМ ПРАВОВОГО РЕГУЛИРОВАНИЯ ЭКОНОМИЧЕСКОЙ ДЕЯТЕЛЬНОСТИ',
                 'СЕКЦИЯ АКТУАЛЬНЫХ ПРОБЛЕМ ПУБЛИЧНОЙ СЛУЖБЫ', 'СЕКЦИЯ ГОСУДАРСТВЕННОГО (КОНСТИТУЦИОННОГО) ПРАВА РОССИИ', 'СЕКЦИЯ ГОСУДАРСТВЕННОГО (КОНСТИТУЦИОННОГО) ПРАВА ЗАРУБЕЖНЫХ СТРАН',
                 'СЕКЦИЯ АДМИНИСТРАТИВНОГО ПРАВА РОССИЙСКОЙ ФЕДЕРАЦИИ', 'СЕКЦИЯ ГРАЖДАНСКОГО ПРАВА', 'СЕКЦИЯ ТРУДОВОГО ПРАВА', 'СЕКЦИЯ СЕМЕЙНОГО ПРАВА', 'СЕКЦИЯ ПРАВА СОЦИАЛЬНОГО ОБЕСПЕЧЕНИЯ',
                 'СЕКЦИЯ НОТАРИАТА В РОССИЙСКОЙ ФЕДЕРАЦИИ', 'СЕКЦИЯ КОРПОРАТИВНОГО ПРАВА', 'СЕКЦИЯ ГРАЖДАНСКОГО И АРБИТРАЖНОГО ПРОЦЕССА', 'СЕКЦИЯ ФИНАНСОВОГО, НАЛОГОВОГО, БАНКОВСКОГО ПРАВА',
                 'СЕКЦИЯ ПРЕДПРИНИМАТЕЛЬСКОГО И КОММЕРЧЕСКОГО ПРАВА', 'СЕКЦИЯ КРИМИНОЛОГИИ', 'СЕКЦИЯ АКТУАЛЬНЫХ ПРОБЛЕМ ОБЩЕЙ ЧАСТИ УГОЛОВНОГО ПРАВА', 'СЕКЦИЯ АКТУАЛЬНЫХ ПРОБЛЕМ ОСОБЕННОЙ ЧАСТИ УГОЛОВНОГО ПРАВА',
                 'СЕКЦИЯ УГОЛОВНОГО ПРАВА И УГОЛОВНОЙ ЮСТИЦИИ', 'СЕКЦИЯ УГОЛОВНОГО ПРОЦЕССА И КРИМИНАЛИСТИКИ', 'СЕКЦИЯ АКТУАЛЬНЫХ ПРОБЛЕМ УГОЛОВНОГО ПРОЦЕССА И КРИМИНАЛИСТИКИ',
                 'СЕКЦИЯ НРАВСТВЕННЫХ ОСНОВ ДЕЯТЕЛЬНОСТИ ЮРИСТА', 'СЕКЦИЯ РУССКОГО ЯЗЫКА', 'СЕКЦИЯ МАССОВОЙ КОММУНИКАЦИИ', 'СЕКЦИЯ СТРАТЕГИЙ И ИННОВАЦИЙ В ИЗДАТЕЛЬСКОМ ДЕЛЕ',
                 'СЕКЦИЯ ЛИТЕРАТУРОВЕДЕНИЯ', 'СЕКЦИЯ ТЕОРИИ И ПРАКТИКИ СВЯЗЕЙ С ОБЩЕСТВЕННОСТЬЮ', 'СЕКЦИЯ ЖУРНАЛИСТИКИ', 'СЕКЦИЯ ЯЗЫКОВЫХ И СОЦИОКУЛЬТУРНЫХ КОМПОНЕНТОВ ДИСКУРСОВ', 'СЕКЦИЯ ЛИНГВОАКСИОЛОГИЧЕСКИХ АСПЕКТОВ ДИСКУРСИВНЫХ ИССЛЕДОВАНИЙ',
                 'СЕКЦИЯ ЛИНГВИСТИЧЕСКИХ АСПЕКТОВ АНГЛОЯЗЫЧНЫХ ДИСКУРСИВНЫХ ПРАКТИК', 'СЕКЦИЯ ПРОБЛЕМ СОВРЕМЕННОЙ АНГЛОЯЗЫЧНОЙ КОММУНИКАЦИИ (НА МАТЕРИАЛЕ РАЗЛИЧНЫХ ТИПОВ ДИСКУРСА)',
                 'СЕКЦИЯ ПРОБЛЕМ ПРАГМАЛИНГВИСТИКИ И ДИСКУРСА В СОВРЕМЕННОМ АНГЛИЙСКОМ ЯЗЫКЕ (ФУНКЦИОНАЛЬНЫЕ АСПЕКТЫ)', 'СЕКЦИЯ МЕТОДИКИ ОБУЧЕНИЯ ИНОСТРАННОМУ ЯЗЫКУ', 'СЕКЦИЯ «ЯЗЫК. КУЛЬТУРА. ОБЩЕСТВО»',
                 'СЕКЦИЯ НЕМЕЦКОЙ ФИЛОЛОГИИ', 'СЕКЦИЯ ОТЕЧЕСТВЕННОЙ ИСТОРИИ И ИСТОРИОГРАФИИ', 'СЕКЦИЯ ИСТОЧНИКОВЕДЕНИЯ И ВСПОМОГАТЕЛЬНЫХ ИСТОРИЧЕСКИХ ДИСЦИПЛИН', 'СЕКЦИЯ «ВЛАСТЬ И ОБЩЕСТВО В РОССИЙСКОЙ ИСТОРИИ»',
                 'СЕКЦИЯ СОЦИОКУЛЬТУРНОГО ПРОСТРАНСТВА РОССИИ В XIX-XX ВВ.', 'СЕКЦИЯ ПРОБЛЕМ РОССИЙСКОЙ ИСТОРИИ', 'СЕКЦИЯ ВСЕОБЩЕЙ ИСТОРИИ', 'СЕКЦИЯ ОРГАНИЗАЦИИ ДОКУМЕНТАЦИОННЫХ ПРОЦЕССОВ', 'СЕКЦИЯ МЕЖДУНАРОДНЫХ ОТНОШЕНИЙ',
                 'СЕКЦИЯ ПОЛИТОЛОГИИ', 'СЕКЦИЯ СОЦИОГУМАНИТАРНОЙ КИБЕРНЕТИКИ', 'СЕКЦИЯ СОЦИОЛОГИИ ТРУДА И УПРАВЛЕНИЯ', 'СЕКЦИЯ ИСТОРИИ СОЦИОЛОГИИ', 'СЕКЦИЯ СОЦИАЛЬНОЙ АНАЛИТИКИ НОВЫХ МЕДИА',
                 'СЕКЦИЯ МЕЖЭТНИЧЕСКИХ ОТНОШЕНИЙ В СОВРЕМЕННОЙ РОССИИ', 'СЕКЦИЯ СОЦИОЛОГИИ СОЦИАЛЬНЫХ ИЗМЕНЕНИЙ', 'СЕКЦИЯ СОЦИОЛОГИИ ПОЛИТИКИ И ГОСУДАРСТВЕННОГО УПРАВЛЕНИЯ', 'СЕКЦИЯ ЭМПИРИЧЕСКОЙ СОЦИОЛОГИИ',
                 'СЕКЦИЯ МЕТОДОЛОГИИ И ПРАКТИКИ СОЦИОЛОГИЧЕСКИХ ИССЛЕДОВАНИЙ', 'СЕКЦИЯ КОНЦЕПТУАЛЬНЫХ ОСНОВ СОЦИАЛЬНОЙ РАБОТЫ', 'СЕКЦИЯ ТЕОРЕТИЧЕСКИХ ПОДХОДОВ К СОЦИОНОМИЧЕСКОЙ ПРАКТИКЕ',
                 'СЕКЦИЯ ОПЫТА СОЦИАЛЬНОЙ РАБОТЫ С РАЗНЫМИ КАТЕГОРИЯМИ ГРАЖДАН', 'СЕКЦИЯ СОЦИАЛЬНЫХ ПРОБЛЕМ СОВРЕМЕННОГО ОБЩЕСТВА', 'СЕКЦИЯ ТЕХНОЛОГИИ СОЦИАЛЬНОЙ РАБОТЫ С РАЗНЫМИ КАТЕГОРИЯМИ ГРАЖДАН',
                 'СЕКЦИЯ ОБЩЕЙ ПСИХОЛОГИИ', 'СЕКЦИЯ СОЦИАЛЬНОЙ ПСИХОЛОГИИ', 'СЕКЦИЯ ПСИХОЛОГИИ РАЗВИТИЯ', 'СЕКЦИЯ ПСИХОЛОГО-ПЕДАГОГИЧЕСКОЙ РАБОТЫ ПО ВОСПИТАНИЮ ДЕТЕЙ И МОЛОДЕЖИ',
                 'СЕКЦИЯ ПСИХОЛОГО-ПЕДАГОГИЧЕСКОГО СОПРОВОЖДЕНИЯ ОБРАЗОВАТЕЛЬНОГО ПРОЦЕССА', 'СЕКЦИЯ АКТУАЛЬНЫХ ПРОБЛЕМ ОБУЧЕНИЯ И ВОСПИТАНИЯ', 'СЕКЦИЯ ПСИХОЛОГО-ПЕДАГОГИЧЕСКОГО ПРОСВЕЩЕНИЯ РОДИТЕЛЕЙ',
                 'СЕКЦИЯ ТЕОРИИ ОБУЧЕНИЯ И ВОСПИТАНИЯ', 'СЕКЦИЯ АКТУАЛЬНЫХ ПРОБЛЕМ ПСИХОЛОГО-ПЕДАГОГИЧЕСКОЙ СЛУЖБЫ В ОБРАЗОВАНИИ', 'СЕКЦИЯ «РИСКИ СОВРЕМЕННОГО ОБРАЗОВАНИЯ»',
                 'СЕКЦИЯ «ЦИФРОВАЯ ОБРАЗОВАТЕЛЬНАЯ СРЕДА»', 'СЕКЦИЯ ИНОСТРАННОГО ЯЗЫКА В ПРОФЕССИОНАЛЬНОЙ КОММУНИКАЦИИ', 'СЕКЦИЯ АНГЛИЙСКОГО ЯЗЫКА В ЮРИСПРУДЕНЦИИ', 'СЕКЦИЯ АНГЛИЙСКОГО ЯЗЫКА В ЕСТЕСТВЕННОНАУЧНОЙ СФЕРЕ',
                 'СЕКЦИЯ НЕМЕЦКОГО ЯЗЫКА И СТРАНОВЕДЕНИЯ ГЕРМАНИИ', 'СЕКЦИЯ АНГЛИЙСКОГО ЯЗЫКА ДЛЯ АКАДЕМИЧЕСКИХ ЦЕЛЕЙ В ЕСТЕСТВЕННОНАУЧНОЙ СФЕРЕ', 'СЕКЦИЯ АНГЛИЙСКОГО ЯЗЫКА ДЛЯ АКАДЕМИЧЕСКИХ ЦЕЛЕЙ В ГУМАНИТАРНОЙ СФЕРЕ',
                 'СЕКЦИЯ «ИНОЯЗЫЧНОЕ ОБРАЗОВАНИЕ»']

# список с индексами секций, содержащих подсекции
sect_with_undersect = [0, 2, 3, 18, 21, 26, 36, 40, 76, 100, 103, 104, 106, 107, 116, 122]

# Создание нового документа Word
doc = docx.Document()

section = doc.sections[0]
data = np.array([], [])
data = table.values
numberOfRecords = len(data)

index_under = int(4)

# Перебор всех секций
for sec in range(0, len(sections_list)):
    # Проверка, содержит ли текущая секция подсекции
    if sec in sect_with_undersect:
        flag_undersect = True
    else:
        flag_undersect = False
    # Добавление заголовка секции
    table = {'Наименование секции': sections_list[sec]}
    head = doc.add_paragraph()
    run = head.add_run(sections_list[sec])
    # Перебор записей в таблице данных
    for numberRecord in range(numberOfRecords):
        if data[numberRecord][2] == sections_list[sec]:
            # Добавление информации о заседании
            if flag_undersect:
                table.update({'Подсекция': data[numberRecord][int(index_under)]})
                if (str(data[numberRecord][int(index_under)])) != 'nan':
                    #run = head.add_run('\n' + data[numberRecord][section])
                    print(data[numberRecord][int(index_under)])
                    head = doc.add_paragraph()
                    print(data[numberRecord][2])
                    run = head.add_run(data[numberRecord][int(index_under)])
                if str(data[numberRecord][20]) != 'nan':
                    table.update({'Номер заседания подсекции': data[numberRecord][20]})
                    run = head.add_run(' (Заседание ' + str(int(data[numberRecord][20])) + ')')
            else:
                if (str(data[numberRecord][3]) != 'nan'):
                    table.update({'Номер заседания секции': int(data[numberRecord][3])})
                    head = doc.add_paragraph()
                    run = head.add_run('Заседание № ' + str(int(data[numberRecord][3])))

            # Создание таблицы для данных о заседании
            tableW = doc.add_table(rows=2, cols=3)

            #tableW.style = 'Table Normal'
            cell = tableW.cell(0, 0)
            cell.text = 'Председатель'
            cell = tableW.cell(0, 1)

            if str(data[numberRecord][25]) != 'nan':
                table.update({'Председатель подсекции': data[numberRecord][25]})
                cell.text = '- ' + data[numberRecord][25]
            elif str(data[numberRecord][21]) != 'nan':
                table.update({'Председатель секции': data[numberRecord][21]})
                cell.text = '- ' + data[numberRecord][21]
            else:
                cell.text = 'Председателя нет..'

            if str(data[numberRecord][26]) != 'nan':
                table.update({'Сопредседатель подсекции': data[numberRecord][26]})
                cell.text += ', сопредседатель - ' + data[numberRecord][26]
            elif str(data[numberRecord][22]) != 'nan':
                table.update({'Сопредседатель секции': data[numberRecord][22]})
                cell.text += ', сопредседатель - ' + data[numberRecord][22]

            cell = tableW.cell(1, 0)
            cell.text = 'Секретарь'
            cell = tableW.cell(1, 1)
            if str(data[numberRecord][27]) != 'nan':
                table.update({'Секретарь подсекции': data[numberRecord][27]})
                cell.text = '- ' + data[numberRecord][27]
            elif str(data[numberRecord][23]) != 'nan':
                table.update({'Секретарь секции': data[numberRecord][23]})
                cell.text = '- ' + data[numberRecord][23]
            else:
                cell.text = 'Секретаря нет'

            if str(data[numberRecord][24]) != 'nan':
                table.update({'Номер телефона секретаря': data[numberRecord][24]})

            cell = tableW.cell(0, 2)
            if str(data[numberRecord][28]) != 'nan':
                table.update({'Дата заседания': data[numberRecord][28]})
                cell.text = 'Заседание '+data[numberRecord][28][9] + ' aпреля'
            else:
                cell.text = 'никогда'

            cell = tableW.cell(1, 2)
            if str(data[numberRecord][29]) != 'nan':
                table.update({'Время': data[numberRecord][29]})
                cell.text = data[numberRecord][29]
            else:
                cell.text = 'нигде и никогда'
            if str(data[numberRecord][31]) != 'nan':
                table.update({'Ссылка на видеоконференцию': data[numberRecord][31]})
                cell.text += ', ' + data[numberRecord][31]
            elif str(data[numberRecord][30]) != 'nan':
                table.update({'Место': data[numberRecord][30]})
                cell.text += ', ' + data[numberRecord][30]
            i = 0
            for speakers in range(32, data.shape[1], 4):
                if str(data[numberRecord][speakers]) != 'nan':
                    i += 1
                else:
                    break

            # Создание таблицы для списка докладчиков
            tableS = doc.add_table(rows=i+1, cols=4)
            # Добавление информации о докладах в таблицу
            i = 0
            tableS.style = 'Table Grid'
            cell = tableS.cell(0, 0)
            cell.text = '№'
            cell = tableS.cell(0, 1)
            cell.text = 'Тема доклада'
            cell = tableS.cell(0, 2)
            cell.text = 'Докладчик'
            cell = tableS.cell(0, 3)
            cell.text = 'Научный руководитель'
            for speakers in range(32, data.shape[1], 4):
                if str(data[numberRecord][speakers]) != 'nan':
                    cell = tableS.cell(i+1, 0)
                    cell.text = str(i+1) + '.'
                    table.update({'Тема доклада [' + str(i) + ']': data[numberRecord][speakers]})
                    cell = tableS.cell(i + 1, 1)
                    cell.text = data[numberRecord][speakers]

                    table.update({'Докладчик[' + str(i) + ']': data[numberRecord][speakers+1]})

                    cell = tableS.cell(i + 1, 2)
                    cell.text = data[numberRecord][speakers+1]
                    table.update({'Научный руководитель[' + str(i) + ']': data[numberRecord][speakers+2]})
                    cell = tableS.cell(i + 1, 3)
                    if str(data[numberRecord][speakers+3]) != 'nan':
                        table.update({'Второй научный руководитель[' + str(i) + ']': data[numberRecord][speakers+3]})
                        cell.text = data[numberRecord][speakers+2] + ',\n' + data[numberRecord][speakers+3]
                    else:
                        cell.text = data[numberRecord][speakers + 2]
                    i+=1
                else:
                    break
            print(table)
    if flag_undersect:
        index_under+=1
        # Сохранение документа
doc.save('test.docx')